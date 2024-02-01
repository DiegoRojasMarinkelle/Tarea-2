import docx
import os
import nltk
import nltk.corpus
import re
import numpy
import spacy
import tkinter as tk

from nltk.tokenize import wordpunct_tokenize
from nltk.stem import PorterStemmer
from nltk.stem import wordnet
from nltk.stem import WordNetLemmatizer
from nltk.corpus import stopwords
from nltk import ne_chunk
from nltk.text import Text
from nltk.util import bigrams
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from googletrans import Translator
from tkinter import filedialog

"""Metodos nltk"""
pst = PorterStemmer()
wnl = WordNetLemmatizer()

"""Metodos Google Translator"""
translator = Translator()

"""Parametros Stop words y puntuación, excluyendo comas y puntos para poder delimitar frases e incisos"""
punctuation = re.compile(r'[-?!;()|]')
stopWords = set(stopwords.words('english'))

"""Traduce las llaves de un diccionario"""
def traducir_Diccionario(kwCategoriesES):
    kwCategoriesEN = dict()
    for k,v in kwCategoriesES.items():
        kwCategoriesEN[translator.translate(k).text] = v

    return kwCategoriesEN

"""Reconoce las clausulas dentro de un documento .docx"""
def reconocer_Clausulas(documento, kwClausulas):
    esPrimeraClausula = True
    esClausula = False
    clausulas = []
    clausulaAux = []
    for paragraph in documento.paragraphs:
        minus = paragraph.text.lower()
        sTokens = wordpunct_tokenize(minus)

        esClausula = False

        """try para parrafos vacios"""
        try:
            fWord = sTokens[0]
            for word in kwClausulas:
                if word in fWord:
                    if esPrimeraClausula == True:
                        clausulaAux.append(minus)
                        esClausula = True
                        esPrimeraClausula = False
                        break
                    else:
                        clausulas.append(clausulaAux)
                        clausulaAux=[]
                        clausulaAux.append(minus)
                        esPrimeraClausula = False
                        esClausula = True
                        break

        except:
            pass

        if esClausula == False and esPrimeraClausula == False:
            clausulaAux.append(minus)

    clausulas.append(clausulaAux)
    return clausulas

"""Traduce una lista de clausulas"""
def traducir_Clausulas(Lista_Clausulas):
    clausulasTraducidas = []
    
    for clausula in Lista_Clausulas:
        parrafoTraducido = []
        for parrafo in clausula:
            parrafoTraducido.append(translator.translate(parrafo).text)
        clausulasTraducidas.append(parrafoTraducido)

    return clausulasTraducidas

"""Remover stop words y puntuación de todas las clausulas"""
def limpiar(Lista_Clausulas_Traducida):

    clausulasPostProcessing = []
    for clausula in Lista_Clausulas_Traducida:
        clausulaPostPunctuation = []
        categoriesAux = []
        
        for parrafo in clausula:
            postPunctuation = []

            """Remover stopwords y traducción de tokens"""
            tokens = wordpunct_tokenize(parrafo)
            filtr = [words.lower() for words in tokens if not words.lower() in stopWords]

            """Remover puntuación"""
            for words in filtr:
                word = punctuation.sub("",words)
                if len(word)>0:
                    postPunctuation.append(word)
            clausulaPostPunctuation.append(postPunctuation)
        clausulasPostProcessing.append(clausulaPostPunctuation)

    return clausulasPostProcessing

"""Realiza el stemming de todas las clausulas a los tokens para la categorización"""
def stemming_Lista(lista_Clausulas_PostLimpieza):
    clausulasPostStemming = []
    for clausula in lista_Clausulas_PostLimpieza:
        clausulaPostStemming = []
        for paragraph in clausula:
            postStemming = []
            for token in paragraph:
                postStemming.append(pst.stem(token))
            clausulaPostStemming.append(postStemming)
        clausulasPostStemming.append(clausulaPostStemming)
   
    return clausulasPostStemming

"""Realiza el stemming de todas las llaves de un diccionario"""
def stemming_Diccionario(diccionario):

    keyWordsStem = []
    categoriesList = []

    for k,v in diccionario.items():
        kTokens = wordpunct_tokenize(k)
        kStemAux = []
        for token in kTokens: 
            kStemAux.append(pst.stem(token))
        keyWordsStem.append(kStemAux)
        categoriesList.append(v)

    return keyWordsStem, categoriesList
                
"""Realiza la categorización de acuerdo a las key words"""
def categorizacion(lista_Clausulas_PostStemming, keyWordsStem, categoriesList):
    categories = []
    for clausula in lista_Clausulas_PostStemming:
        categoriesAux = []
        
        for paragraph in clausula:
            bigramsParagraph = list(nltk.bigrams(paragraph))
            
            for stem in keyWordsStem:
                """¿Que pasa si hay trigrams? ¿condición mayor que 3?"""
                bigramCat = categoriesList[keyWordsStem.index(stem)]
                
                if bigramCat not in categoriesAux:
                    if len(stem)>1:
                        for bigram in stem:
                            if bigram in bigramsParagraph:
                                categoriesAux.append(bigramCat)

                        for bigram in list(nltk.bigrams(list(reversed(stem)))):
                            if bigram in bigramsParagraph:
                                categoriesAux.append(bigramCat)
                                    
                    else:
                        if stem[0] in paragraph:
                            categoriesAux.append(bigramCat)

        categories.append(categoriesAux)
        
    return categories

"""Lemmatization de los tokens posteriores a limpieza de stopWords y puntuación"""
def lemmatize(lista_Limpia):
    lemmatized = []
    parrafo_Tagged = []
    
    for clausula in lista_Limpia:
        paragraph_Lemmatized = []
        for parrafo in clausula:
            parrafo_Tagged = nltk.pos_tag(parrafo)
            token_Lemmatized = []
            for token in parrafo_Tagged:
                """Es posible indicar el tipo de Part Of Speech que es la palabra"""
                if token[0] == "parties" or token[1][:2] == "VB":
                    token_Lemmatized.append(token[0])
                elif token[0] == "ani":
                    token_Lemmatized.append("ANI")
                elif token[0] == "covid":
                    token_Lemmatized.append("Covid")
                else:
                    token_Lemmatized.append(wnl.lemmatize(token[0], pos = "v"))
            paragraph_Lemmatized.append(token_Lemmatized)
        lemmatized.append(paragraph_Lemmatized)
    return lemmatized

"""Realiza el POS tagging (Parts Of Speech) para los parrafos de cada una de las clausulas"""
def POS_Tags(lista_Lemmatized):
    POS_Tagged = []
    
    for clausula in lista_Lemmatized:
        paragraph_Tagged = []
        for parrafo in clausula:
            paragraph_Tagged.append(nltk.pos_tag(parrafo))
        POS_Tagged.append(paragraph_Tagged)
    
    return POS_Tagged

def NER(Lista_POS_Tags):

    """Definición de lista de Name Entity"""
    NER = ["concessionaire", "concessionaires", "parties", "party", "ANI", "ani"]
    NN = ["covid", "Covid"]

    """Definición de listas"""
    lista_NER = []

    for clausula in Lista_POS_Tags:
        parrafo_NER = []

        for parrafo in clausula:
            tupla_NER = []

            for token, tag in parrafo:
                if token in NER:
                    tupla_NER.append((token, "NNP"))
                elif token in NN:
                    tupla_NER.append((token, "NN"))
                else:
                    tupla_NER.append((token, tag))
            parrafo_NER.append(tupla_NER)
        lista_NER.append(parrafo_NER)

    return lista_NER

"""Extrae la frase con la idea principal de cada clausula"""
def extractingSentence(Lista_NER, Lista_POS_Tags):
    sentence_List_Tagged = []
    sentence_List_NER = []

    for clausula in Lista_NER:
        """Temporales de posiciones"""
        pos_Puntos = []
        pos_Comas = []
        pos_VB = []
        pos_ENT = []
        sentence_Tagged = []
        sentence_NER = []

        for idx, token in enumerate(clausula[0]):
            if token[1] == "." or token[1] == ":":
                pos_Puntos.append(idx)
            elif token[1][:2] == "VB":
                pos_VB.append(idx)
            elif token[1] == "NNP":
                pos_ENT.append(idx)

        for idx_Puntos in (range(len(pos_Puntos))):
            if len(pos_ENT) == 0:
                sentence_NER = clausula
                sentence_Tagged = clausula
                break
            if len(sentence_Tagged) > 0:
                break
            rango = range(pos_Puntos[idx_Puntos], pos_Puntos[idx_Puntos + 1])
            for idx_VB in range(len(pos_VB)):

                for idx_ENT in range(len(pos_ENT)):
                    if pos_ENT[idx_ENT] > pos_Puntos[idx_Puntos + 1]:
                        break
                    if pos_VB[idx_VB] in rango and pos_ENT[idx_ENT] in rango:
                        sentence_Tagged.extend(Lista_POS_Tags[Lista_NER.index(clausula)][0][pos_Puntos[idx_Puntos] + 1:pos_Puntos[idx_Puntos + 1]])
                        sentence_List_Tagged.append(sentence_Tagged)
                        sentence_NER.extend(clausula[0][pos_Puntos[idx_Puntos] + 1:pos_Puntos[idx_Puntos + 1]])
                        sentence_List_NER.append(sentence_NER)
                        break

                if pos_VB[idx_VB] > pos_Puntos[idx_Puntos + 1]:
                    break
                if len(sentence_Tagged) > 0:
                    break

    return sentence_List_NER, sentence_List_Tagged

"""Remover incisos dentro de las frases claves"""
def remover_Incisos(lista_NER, lista_Sentences):

    sentence = []

    for clausula in lista_NER:
        pos_Comas = []
        pos_ENT = []
        rangos_Remover = []
        c = lista_Sentences[lista_NER.index(clausula)]

        for idx, token in enumerate(clausula):
            if token[1] == "NNP":
                pos_ENT.append(idx)
            elif token[1] == ",":
                pos_Comas.append(idx)

        for idx_Comas in range(len(pos_Comas)):
            if len(pos_ENT) == 0:
                break

            for idx_ENT in range(len(pos_ENT)):
                if idx_Comas == 0:
                    if pos_ENT[idx_ENT] < pos_Comas[idx_Comas]:
                        break
                    elif idx_ENT == len(pos_ENT) - 1:
                        rangos_Remover.append(range(0,pos_Comas[idx_Comas]))
                else:
                    if pos_ENT[idx_ENT] in range(pos_Comas[idx_Comas - 1], pos_Comas[idx_Comas]):
                        break
                    elif pos_ENT[idx_ENT] > pos_Comas[idx_Comas]:
                        rangos_Remover.append(range(pos_Comas[idx_Comas - 1], pos_Comas[idx_Comas]))
                        break
                    elif idx_ENT == len(pos_ENT) - 1:
                        rangos_Remover.append(range(pos_Comas[idx_Comas - 1], pos_Comas[idx_Comas]))

        for r_Index in range(len(rangos_Remover)):
            lim_Inferior = rangos_Remover[len(rangos_Remover) - r_Index - 1][0]
            lim_Superior = rangos_Remover[len(rangos_Remover) - r_Index - 1][-1]
            del lista_Sentences[lista_NER.index(clausula)][lim_Inferior: lim_Superior + 1]

    return lista_Sentences

"""Realiza el chunking, de acuerdo a cada una de las reglas definidas, para los parrafos de cada una de las clausulas"""
def chunking(lista_Sentences):
    lista_Chunked = []

    chunk_Grammar_Subject = r"""Subject:{<NNP>}"""

    chunk_Grammar_Verb = r"""Verb:{<VB.?>*}"""

    chunk_Grammar_Object = r"""Object:{<.*>*<NNP>?}"""

    chunk_Grammar_SVO_1 = r"""SVO:
                                {<NNP><.*>*<NNP>}
                                }<JJ>{
                            """
    chunk_Grammar_SVO_2 = r"""SVO:
                                {<NNP><.*>*}
                                }<JJ>{
                            """

    for sentence in lista_Sentences:

        parrafo_Chunked = []

        sentence_Part_S = []
        sentence_Part_O = []
        sentence_SVO = []
        structure = ""

        for token in sentence:
            c_Subjects = 0
            repetido_Subject = False
            if token[1] == "NNP":
                structure = structure + "s"
            elif token[1][:2] == "VB":
                structure = structure + "v"
            if not structure == "":
                if structure[0] == "s":
                    previous_Subject = ""
                    for token in sentence:
                        if not token[0] == ",":
                            if token[1] == "NNP":
                                if wnl.lemmatize(token[0]) == wnl.lemmatize(previous_Subject):
                                    repetido_Subject = True
                                    break
                                else:
                                    c_Subjects += 1
                                    previous_Subject = token[0]
                            sentence_SVO.append(token)
                    break
                elif token[1] == "," and structure[0] == "v":
                    sentence_Part_O = sentence[:sentence.index(token)]
                    sentence_Part_S = sentence[sentence.index(token) + 1:]
                    break

        if structure[0] == "v" and not sentence_Part_S and not sentence_Part_O and "s" in structure:
            structure = "s"
            for token in sentence:
                if not token[0] == ",":
                    if token[1] == "NNP":
                        if wnl.lemmatize(token[0]) == wnl.lemmatize(previous_Subject):
                            break
                        else:
                            c_Subjects += 1
                            previous_Subject = token[0]
                    sentence_SVO.append(token)

        elif structure[0] == "v" and not sentence_Part_S and not sentence_Part_O and "s" not in structure:
            structure = "n"

        if structure[0] == "s":
            if c_Subjects == 2 and repetido_Subject == True:
                chunkParserSVO_1 = nltk.RegexpParser(chunk_Grammar_SVO_1)
                chunkedSVO_1 = chunkParserSVO_1.parse(sentence_SVO)
                for subtree in chunkedSVO_1.subtrees(lambda t: t.label() == "SVO"):
                        for leaf in subtree.leaves():
                            parrafo_Chunked.append(leaf[0] + " " + leaf[1])
            else:
                chunkParserSVO_2 = nltk.RegexpParser(chunk_Grammar_SVO_2)
                chunkedSVO_2 = chunkParserSVO_2.parse(sentence_SVO)
                for subtree in chunkedSVO_2.subtrees(lambda t: t.label() == "SVO"):
                    for leaf in subtree.leaves():
                        parrafo_Chunked.append(leaf[0] + " " + leaf[1])

        elif structure[0] == "v":
                chunkParserSubject = nltk.RegexpParser(chunk_Grammar_Subject)
                chunkedSubject = chunkParserSubject.parse(sentence_Part_S)

                chunkParserVerb = nltk.RegexpParser(chunk_Grammar_Verb)

                chunkParserObject = nltk.RegexpParser(chunk_Grammar_Object)
                chunkedObject = chunkParserObject.parse(sentence_Part_O)

                for subtree in chunkedSubject.subtrees(filter=lambda t: t.label() == "Subject"):
                        for leaf in subtree.leaves():
                            parrafo_Chunked.append(leaf[0] + " " + leaf[1])


                if chunkParserVerb.parse(sentence_Part_S).subtrees(filter=lambda t: t.label() == "Verb"):
                    chunkedVerb = chunkParserVerb.parse(sentence_Part_S)

                    for subtree in chunkedVerb.subtrees(filter=lambda t: t.label() == "Verb"):
                        for leaf in subtree.leaves():
                            parrafo_Chunked.append(leaf[0] + " " + leaf[1])


                for subtree in chunkedObject.subtrees(filter=lambda t: t.label() == "Object"):
                        for leaf in subtree.leaves():
                            parrafo_Chunked.append(leaf[0] + " " + leaf[1])

        elif structure == "n":
            chunkParserVerb = nltk.RegexpParser(chunk_Grammar_Verb)
            chunkedObject = chunkParserObject.parse(sentence)
            chunkParserObject = nltk.RegexpParser(chunk_Grammar_Object)
            chunkedObject = chunkParserObject.parse(sentence)

            for subtree in chunkedSubject.subtrees(filter=lambda t: t.label() == "Subject"):
                    for leaf in subtree.leaves():
                        parrafo_Chunked.append(leaf[0] + " " + leaf[1])

            for subtree in chunkedObject.subtrees(filter=lambda t: t.label() == "Object"):
                    for leaf in subtree.leaves():
                        parrafo_Chunked.append(leaf[0] + " " + leaf[1])

        lista_Chunked.append(parrafo_Chunked)

    return lista_Chunked

"""Exporta una lista de clausulas a un documento .docx"""
def exportar_Docx(lista_Exportar, lista_Categorias, file_Name):
    """Definicion Parametros documento .docx"""
    file = Document()
    style = file.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(15)
    font.color.rgb = RGBColor(51, 0, 0)

    """Titulo documento"""
    file.add_paragraph("Clausulas " + file_Name[:file_Name.index(".")])
    
    contClausulas = 1
    try:
        
        for clausula in lista_Exportar:
            file.add_heading("Clausula "+ str(contClausulas) + ", Category:" + str(lista_Categorias[lista_Exportar.index(clausula)]))
            contClausulas = contClausulas + 1
            file.add_paragraph(" ".join(clausula))

        file.save("C:/Users/asus/Desktop/" + "Resultados NLP " + file_Name)

    except Exception as ex:
        formato = "Excepción de tipo {0} occurrio. Argumentos:\n{1!r}"
        mensaje = formato.format(type(ex).__name__, ex.args)
        print(mensaje)

def main():

    """importar documento"""
    root = tk.Tk()
    root.withdraw()

    file_Path = filedialog.askopenfilename()

    pos_Segmentation = []
    for idx, character in enumerate(file_Path):
        if character == "/":
            pos_Segmentation.append(idx)

    file_Name = file_Path[pos_Segmentation[len(pos_Segmentation)-1] + 1:]

    doc = docx.Document(file_Path)

    """Definir parametros,variables y arreglos"""
    kwClausulas=["primer", "segund", "tercer", "cuart", "quint",
                 "sext", "septim", "octav", "noven", "decim", "undecim"]

    """Definición y traducción de lexicon"""
    kwCategoriesES = dict()
    kwCategoriesES = {"compensación recursos":"Compensation",
                      "exención":"Cat2",
                      "partes":"Cat3",
                      "afectación ingresos":"compensation",
                      "impacto ingresos":"compensation"}

    """llamar funcion traducir_Diccionario"""
    kwCategoriesEN = traducir_Diccionario(kwCategoriesES)
    """llamar funcion reconocer_Clausulas"""
    clausulas = reconocer_Clausulas(doc, kwClausulas)
    """Llamar función traducir_Clausulas"""
    clausulasEN = traducir_Clausulas(clausulas)
    """Llamar función limpiar"""
    clean = limpiar(clausulasEN)
    """Llamar función stemming_lista"""
    stemClausulas = stemming_Lista(clean)
    """Llamar función steming_diccionario"""
    keyWordsStem, categoriesList = stemming_Diccionario(kwCategoriesEN)
    """Llamar función categorización"""
    categorizado = categorizacion(stemClausulas, keyWordsStem, categoriesList)
    """Llamar función lemmatized"""
    lemmatized = lemmatize(clean)
    """Llamar función POS_Tags"""
    tagged = POS_Tags(lemmatized)
    """Llamar función Name Entity Recognition"""
    entity_Recognition = NER(tagged)
    """Llamar función extractingSentence"""
    sentences_NER, sentences_Tagged = extractingSentence(entity_Recognition, tagged)
    """Remover incisos"""
    incisos = remover_Incisos(sentences_NER, sentences_NER)
    """Llamar función chunking"""
    chunked = chunking(incisos)
    """Llamar función exportarDocx"""
    exportar_Docx(chunked, categorizado, file_Name)

main()
